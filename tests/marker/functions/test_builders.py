import pytest
from docx import Document as DocxDocument

from marker.domain.models import ProjectMetadata, ReportStructure
from marker.functions.builders import (
    build_metadata,
    build_skeleton,
    build_skeleton_bytes,
)


def _make_structure(*headings: str) -> ReportStructure:
    return ReportStructure.model_validate(
        {"Sections": [{"Heading": h} for h in headings]}
    )


def _make_metadata(
    *,
    project_number: str = "P-001",
    client_name: str = "Acme Corp",
    site_address: str = "123 Main St, Auckland",
    date: str = "2026-06-05",
) -> ProjectMetadata:
    return ProjectMetadata.model_validate(
        {
            "Project Number": project_number,
            "Client Name": client_name,
            "Site Address": site_address,
            "Date": date,
        }
    )


def _headings_from_docx(path) -> list[str]:
    doc = DocxDocument(str(path))
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


class TestBuildSkeleton:
    def test_writes_docx_file(self, tmp_path) -> None:
        structure = _make_structure("Introduction", "Methodology", "Conclusion")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        assert output.exists()

    def test_headings_written_in_order(self, tmp_path) -> None:
        structure = _make_structure("Introduction", "Methodology", "Conclusion")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        assert _headings_from_docx(output) == [
            "Introduction",
            "Methodology",
            "Conclusion",
        ]

    def test_single_section(self, tmp_path) -> None:
        structure = _make_structure("Executive Summary")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        assert _headings_from_docx(output) == ["Executive Summary"]

    def test_raises_on_empty_structure(self) -> None:
        with pytest.raises(Exception, match="empty"):
            ReportStructure.model_validate({"Sections": []})

    def test_sections_with_special_characters(self, tmp_path) -> None:
        structure = _make_structure("Site & Location", "Results (Preliminary)")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        assert _headings_from_docx(output) == [
            "Site & Location",
            "Results (Preliminary)",
        ]


class TestBuildMetadata:
    def test_metadata_table_present_in_docx(self, tmp_path) -> None:
        structure = _make_structure("Introduction")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        doc = DocxDocument(str(output))
        assert len(doc.tables) == 1

    def test_metadata_table_has_four_rows(self, tmp_path) -> None:
        structure = _make_structure("Introduction")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        doc = DocxDocument(str(output))
        assert len(doc.tables[0].rows) == 4

    def test_metadata_table_labels(self, tmp_path) -> None:
        structure = _make_structure("Introduction")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        doc = DocxDocument(str(output))
        table = doc.tables[0]
        labels = [table.cell(i, 0).text for i in range(4)]
        assert labels == ["Project Number", "Client Name", "Site Address", "Date"]

    def test_metadata_table_values(self, tmp_path) -> None:
        structure = _make_structure("Introduction")
        metadata = _make_metadata(
            project_number="GEO-042",
            client_name="Kiwi Construction",
            site_address="99 Queen St, Wellington",
            date="2026-01-15",
        )
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        doc = DocxDocument(str(output))
        table = doc.tables[0]
        assert table.cell(0, 1).text == "GEO-042"
        assert table.cell(1, 1).text == "Kiwi Construction"
        assert table.cell(2, 1).text == "99 Queen St, Wellington"
        assert table.cell(3, 1).text == "2026-01-15"

    def test_metadata_table_before_sections(self, tmp_path) -> None:
        structure = _make_structure("Introduction", "Conclusion")
        metadata = _make_metadata()
        output = tmp_path / "report.docx"

        build_skeleton(structure, metadata, output)

        doc = DocxDocument(str(output))
        body_elements = doc.element.body
        table_idx = next(
            i for i, el in enumerate(body_elements) if el.tag.endswith("}tbl")
        )
        heading_idx = next(
            i
            for i, el in enumerate(body_elements)
            if el.tag.endswith("}p")
            and any(
                r.tag.endswith("}pStyle")
                and r.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                    "",
                ).startswith("Heading")
                for r in el.iter()
            )
        )
        assert table_idx < heading_idx

    def test_build_metadata_recording_facade(self, recording_facade) -> None:
        metadata = _make_metadata()

        build_metadata(metadata, recording_facade)

        assert recording_facade.tables == [(4, 2)]
        assert recording_facade.cell_writes[(0, 0, 0)] == "Project Number"
        assert recording_facade.cell_writes[(0, 0, 1)] == "P-001"
        assert recording_facade.cell_writes[(0, 1, 0)] == "Client Name"
        assert recording_facade.cell_writes[(0, 1, 1)] == "Acme Corp"
        assert recording_facade.cell_writes[(0, 2, 0)] == "Site Address"
        assert recording_facade.cell_writes[(0, 2, 1)] == "123 Main St, Auckland"
        assert recording_facade.cell_writes[(0, 3, 0)] == "Date"
        assert recording_facade.cell_writes[(0, 3, 1)] == "2026-06-05"

    def test_date_formatted_as_iso(self, recording_facade) -> None:
        metadata = _make_metadata(date="2026-01-07")

        build_metadata(metadata, recording_facade)

        assert recording_facade.cell_writes[(0, 3, 1)] == "2026-01-07"


class TestBuildSkeletonBytes:
    def test_build_skeleton_bytes_returns_docx_magic(self) -> None:
        structure = _make_structure("Introduction", "Conclusions")
        metadata = _make_metadata()

        result = build_skeleton_bytes(structure, metadata)

        assert len(result) > 0
        assert result[:4] == b"PK\x03\x04"
