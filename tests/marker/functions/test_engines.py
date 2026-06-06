from docx import Document as DocxDocument

from marker.functions.engines import PythonDocxFacade


class TestPythonDocxFacade:
    def test_writes_heading_to_docx(self, tmp_path) -> None:
        facade = PythonDocxFacade()
        facade.add_heading("Introduction", level=1)
        output = tmp_path / "test.docx"
        facade.save(str(output))

        doc = DocxDocument(str(output))
        texts = [p.text for p in doc.paragraphs if p.text]
        assert "Introduction" in texts

    def test_writes_table_to_docx(self, tmp_path) -> None:
        facade = PythonDocxFacade()
        facade.add_table(rows=2, cols=3)
        output = tmp_path / "test.docx"
        facade.save(str(output))

        doc = DocxDocument(str(output))
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 2
        assert len(doc.tables[0].columns) == 3

    def test_add_table_returns_index(self) -> None:
        facade = PythonDocxFacade()
        first = facade.add_table(rows=2, cols=2)
        second = facade.add_table(rows=3, cols=3)

        assert first == 0
        assert second == 1

    def test_write_table_cell_stores_text(self, tmp_path) -> None:
        facade = PythonDocxFacade()
        idx = facade.add_table(rows=2, cols=2)
        facade.write_table_cell(idx, 0, 0, "Label")
        facade.write_table_cell(idx, 0, 1, "Value")
        output = tmp_path / "cells.docx"
        facade.save(str(output))

        doc = DocxDocument(str(output))
        assert doc.tables[0].cell(0, 0).text == "Label"
        assert doc.tables[0].cell(0, 1).text == "Value"

    def test_writes_heading_and_table(self, tmp_path) -> None:
        facade = PythonDocxFacade()
        facade.add_heading("Site Investigation Report", level=1)
        facade.add_table(rows=4, cols=2)
        output = tmp_path / "combined.docx"
        facade.save(str(output))

        doc = DocxDocument(str(output))
        texts = [p.text for p in doc.paragraphs if p.text]
        assert "Site Investigation Report" in texts
        assert len(doc.tables) == 1
