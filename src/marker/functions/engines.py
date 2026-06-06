"""Concrete document engine implementations."""

from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table as DocxTable


class PythonDocxFacade:
    """python-docx backed implementation of DocumentFacade.

    Args:
        template: Path to an existing .docx template file. When omitted,
            a blank document is created.
    """

    def __init__(self, template: Path | None = None) -> None:
        if template is not None:
            self._doc = DocxDocument(str(template))
        else:
            self._doc = DocxDocument()
        self._tables: list[DocxTable] = []

    def add_heading(self, text: str, level: int) -> None:
        """Add a heading paragraph to the document.

        Args:
            text: Heading text content.
            level: Heading level (1 = top-level).
        """
        self._doc.add_heading(text, level=level)

    def add_table(self, rows: int, cols: int) -> int:
        """Insert an empty table and return its zero-based index.

        Args:
            rows: Number of rows.
            cols: Number of columns.

        Returns:
            Zero-based index of the newly added table.
        """
        table = self._doc.add_table(rows=rows, cols=cols)
        self._tables.append(table)
        return len(self._tables) - 1

    def write_table_cell(self, table_index: int, row: int, col: int, text: str) -> None:
        """Write text into a specific table cell.

        Args:
            table_index: Zero-based index of the table (returned by add_table).
            row: Zero-based row index.
            col: Zero-based column index.
            text: Cell content.
        """
        self._tables[table_index].cell(row, col).text = text

    def save(self, path: str) -> None:
        """Write the document to disk.

        Args:
            path: Destination file path.
        """
        self._doc.save(path)
